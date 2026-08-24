#!/usr/bin/env python3
"""
Generate the reproducible research update:

    reports/baseline_data_modelling_update.docx

Project:
    Urban Green Cover Thresholds for PM2.5 Mitigation:
    A Spatial Causal Machine Learning Framework for Delhi NCR

Purpose:
    Build a human-readable, research-style Update 2 report from the frozen
    V2 modeling workflow. The script does not alter datasets. It reads the
    final V2 panel/split artifacts when available, reuses existing modeling
    visualizations from the repository, and embeds validated baseline results.

Expected repository structure (relative to project root):
    data/ml_ready/master_modeling_dataset.csv
    data/ml_ready/master_modeling_dataset_v2.csv
    data/modeling_final/train.csv
    data/modeling_final/test.csv
    data/modeling_final/split_manifest.csv
    data/modeling_final/validation_report.csv
    data/modeling_final/distribution_diagnostics.csv
    data/03_features/feat_era5_met.csv
    data/03_features/feat_worldpop.csv
    data/03_features/feat_osm_roads.csv
    data/03_features/feat_worldcover.csv
    reports/figures/modeling/*.png

The script is intentionally tolerant of alternative figure filenames: it
searches by keywords and uses the first matching PNG for each figure slot.
"""

from __future__ import annotations

import argparse
import json
import math
from pathlib import Path
from typing import Dict, Iterable, List, Optional

import pandas as pd
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


TITLE = (
    "Urban Green Cover Thresholds for PM₂.₅ Mitigation: "
    "A Spatial Causal Machine Learning Framework for Delhi NCR"
)
REPORT_TITLE = "Update 2 — Baseline Modelling and Resplitting of Data"
ANALYTICAL_PERIOD = "2022–2025"
TARGET = "pm25"
SEED = 42

# Current frozen results supplied in the research update. These are used as a
# fallback when the corresponding CSVs are not present locally.
MODEL_RESULTS = [
    {
        "model": "LightGBM",
        "train_r2": 0.998916844,
        "test_r2": 0.897481534,
        "test_rmse": 20.378312853,
        "test_mae": 10.460090407,
        "test_median_ae": 6.512762713,
        "cv_r2_mean": 0.946998828,
        "cv_r2_std": 0.0054277678,
        "cv_rmse_mean": 15.633632564,
        "cv_mae_mean": 10.26580353,
        "r2_gap": 0.10143531,
    },
    {
        "model": "Random Forest",
        "train_r2": 0.990700661,
        "test_r2": 0.875905802,
        "test_rmse": 22.420377114,
        "test_mae": 12.24346966,
        "test_median_ae": 7.50042,
        "cv_r2_mean": 0.924150568,
        "cv_r2_std": 0.008007538,
        "cv_rmse_mean": 18.692022643,
        "cv_mae_mean": 12.073575334,
        "r2_gap": 0.114794859,
    },
    {
        "model": "Linear Regression",
        "train_r2": 0.876668293,
        "test_r2": -1.7145525,
        "test_rmse": 104.861555929,
        "test_mae": 26.275765168,
        "test_median_ae": 16.134595448,
        "cv_r2_mean": 0.814697108,
        "cv_r2_std": 0.029494744,
        "cv_rmse_mean": 29.069350457,
        "cv_mae_mean": 22.011290178,
        "r2_gap": 2.591220793,
    },
    {
        "model": "Ridge Regression",
        "train_r2": 0.870733668,
        "test_r2": -2.129228645,
        "test_rmse": 112.58637168,
        "test_mae": 26.755592223,
        "test_median_ae": 15.665263146,
        "cv_r2_mean": 0.819723672,
        "cv_r2_std": 0.0191801246,
        "cv_rmse_mean": 28.737494184,
        "cv_mae_mean": 21.479841612,
        "r2_gap": 2.999962313,
    },
]

YEARWISE = {
    "Linear Regression": [(-12.855, 2022), (0.903, 2023), (0.680, 2024), (0.821, 2025)],
    "Ridge Regression": [(-15.055, 2022), (0.894, 2023), (0.650, 2024), (0.826, 2025)],
    "Random Forest": [(0.880, 2022), (0.948, 2023), (0.718, 2024), (0.940, 2025)],
    "LightGBM": [(0.914, 2022), (0.963, 2023), (0.740, 2024), (0.961, 2025)],
}

SEASONWISE = {
    "LightGBM": {"Winter": 0.874342252, "Summer": 0.765983268, "Monsoon": 0.529158425, "Post-monsoon": 0.620354358},
    "Random Forest": {"Winter": 0.818773906, "Summer": 0.677110473, "Monsoon": 0.443524640, "Post-monsoon": 0.573100440},
    "Linear Regression": {"Winter": 0.684554249, "Summer": 0.290678598, "Monsoon": -309.493344116, "Post-monsoon": 0.535242228},
    "Ridge Regression": {"Winter": 0.641886238, "Summer": 0.303950572, "Monsoon": -358.458386398, "Post-monsoon": 0.497339797},
}

GROUPED_RESULTS = [
    ("Linear Regression", -19.852880109, 0.668802285),
    ("Ridge Regression", -3.056857753, 0.716245176),
    ("Random Forest", 0.867214147, 0.856382435),
    ("LightGBM", 0.864342121, 0.872892896),
]

FEATURE_GROUPS = {
    "Green cover": {"Linear Regression": 13.429332, "Ridge Regression": 25.34272, "Random Forest": 4.969528, "LightGBM": 46.055016},
    "Meteorology": {"Linear Regression": 22.378509, "Ridge Regression": 18.049943, "Random Forest": 17.325999, "LightGBM": 23.068687},
    "Pollution / anthropogenic": {"Linear Regression": 32.483338, "Ridge Regression": 13.986067, "Random Forest": 1.29624, "LightGBM": 13.473892},
    "Population": {"Linear Regression": 3.960801, "Ridge Regression": 2.565284, "Random Forest": 0.370474, "LightGBM": 0.601219},
    "Road infrastructure": {"Linear Regression": 2.386078, "Ridge Regression": 5.187542, "Random Forest": 1.109386, "LightGBM": 2.643716},
    "Land cover": {"Linear Regression": 12.166356, "Ridge Regression": 16.853458, "Random Forest": 1.420641, "LightGBM": 7.42876},
    "Spatial / temporal": {"Linear Regression": 7.193285, "Ridge Regression": 11.41541, "Random Forest": 72.58266, "LightGBM": 4.266183},
}

RESIDUALS = {
    "mean": -1.708,
    "median": -0.361,
    "std": 20.338,
    "skewness": -6.837,
    "max_underprediction": 44.641,
    "max_overprediction": -262.726,
}


# ---------- Utility functions ----------

def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_text(cell, text: str, bold: bool = False, color: str = "000000", size: int = 9) -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(str(text))
    run.bold = bold
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def format_table(table, header_fill="243447") -> None:
    if not table.rows:
        return
    set_repeat_table_header(table.rows[0])
    for c in table.rows[0].cells:
        set_cell_shading(c, header_fill)
        for p in c.paragraphs:
            for r in p.runs:
                r.font.color.rgb = RGBColor(255, 255, 255)
                r.bold = True
                r.font.size = Pt(8.5)
    for row in table.rows[1:]:
        for c in row.cells:
            for p in c.paragraphs:
                for r in p.runs:
                    r.font.size = Pt(8.5)


def add_table(doc: Document, headers: List[str], rows: Iterable[Iterable[object]], widths: Optional[List[float]] = None) -> None:
    data = [list(r) for r in rows]
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for j, h in enumerate(headers):
        set_cell_text(table.rows[0].cells[j], h, bold=True, color="FFFFFF", size=8.5)
        set_cell_shading(table.rows[0].cells[j], "243447")
    for row_data in data:
        cells = table.add_row().cells
        for j, value in enumerate(row_data):
            set_cell_text(cells[j], value, size=8.5)
    format_table(table)
    doc.add_paragraph()


def add_bullets(doc: Document, items: Iterable[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(2)
        p.add_run(item)


def add_caption(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(9)
    r = p.add_run(text)
    r.italic = True
    r.font.size = Pt(9)
    r.font.color.rgb = RGBColor(90, 90, 90)


def figure_index(paths: Iterable[Path]) -> Dict[str, Path]:
    return {p.name.lower(): p for p in paths}


def find_figure(root: Path, keywords: Iterable[str], used: Optional[set[str]] = None) -> Optional[Path]:
    search_dirs = [
        root / "reports" / "figures" / "modeling",
        root / "reports" / "figures",
        root / "data" / "modeling_final",
        root / "data" / "modeling" / "results",
    ]
    candidates: List[Path] = []
    for d in search_dirs:
        if d.exists():
            candidates.extend(sorted(d.glob("*.png")))
    keyset = [k.lower() for k in keywords]
    scored = []
    for p in candidates:
        if used is not None and p.resolve().as_posix() in used:
            continue
        name = p.name.lower()
        score = sum(1 for k in keyset if k in name)
        if score:
            scored.append((score, len(name), p))
    if scored:
        scored.sort(key=lambda x: (-x[0], x[1], x[2].name))
        return scored[0][2]
    return None


def add_figure(doc: Document, fig: Optional[Path], caption: str, placeholder_text: str) -> None:
    if fig and fig.exists():
        try:
            doc.add_picture(str(fig), width=Inches(6.4))
            doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
            add_caption(doc, caption)
            return
        except Exception as exc:  # pragma: no cover
            p = doc.add_paragraph(f"Figure could not be embedded: {fig.name} ({exc})")
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    else:
        p = doc.add_paragraph(placeholder_text)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.runs[0].italic = True
        add_caption(doc, caption + " (figure not found at generation time)")


def fmt(v: object, digits: int = 3) -> str:
    if isinstance(v, float):
        return f"{v:.{digits}f}"
    return str(v)


def load_if_exists(path: Path) -> Optional[pd.DataFrame]:
    return pd.read_csv(path) if path.exists() else None


def get_model_results(root: Path) -> List[dict]:
    candidates = [
        root / "data" / "modeling" / "results" / "model_comparison.csv",
        root / "data" / "modeling_final" / "model_comparison.csv",
        root / "data" / "modeling_final" / "results" / "model_comparison.csv",
    ]
    df = next((load_if_exists(p) for p in candidates if p.exists()), None)
    if df is None:
        return MODEL_RESULTS
    rows = []
    for _, r in df.iterrows():
        rows.append({
            "model": r.get("Model", r.get("model")),
            "train_r2": float(r.get("Train R²", r.get("train_r2"))),
            "test_r2": float(r.get("Test R²", r.get("test_r2"))),
            "test_rmse": float(r.get("Test RMSE", r.get("test_rmse"))),
            "test_mae": float(r.get("Test MAE", r.get("test_mae"))),
            "test_median_ae": float(r.get("Test MedianAE", r.get("test_median_ae"))),
            "cv_r2_mean": float(r.get("CV R² mean", r.get("cv_r2_mean"))),
            "cv_r2_std": float(r.get("CV R² std", r.get("cv_r2_std"))),
            "cv_rmse_mean": float(r.get("CV RMSE mean", r.get("cv_rmse_mean"))),
            "cv_mae_mean": float(r.get("CV MAE mean", r.get("cv_mae_mean"))),
            "r2_gap": float(r.get("R² gap", r.get("r2_gap"))),
        })
    return rows


def add_header_footer(doc: Document) -> None:
    section = doc.sections[0]
    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    rr = header.add_run("Baseline Modelling Update 2 | Delhi NCR PM₂.₅")
    rr.font.size = Pt(8)
    rr.font.color.rgb = RGBColor(110, 110, 110)

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rr = footer.add_run("Reproducible research record | Analytical period 2022–2025")
    rr.font.size = Pt(8)
    rr.font.color.rgb = RGBColor(110, 110, 110)


def build_document(root: Path, output: Path) -> None:
    output.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()
    used_figures: set[str] = set()
    sec = doc.sections[0]
    sec.top_margin = Inches(0.68)
    sec.bottom_margin = Inches(0.68)
    sec.left_margin = Inches(0.75)
    sec.right_margin = Inches(0.75)

    styles = doc.styles
    styles["Normal"].font.name = "Aptos"
    styles["Normal"].font.size = Pt(10.2)
    styles["Title"].font.name = "Aptos Display"
    styles["Heading 1"].font.name = "Aptos Display"
    styles["Heading 2"].font.name = "Aptos"
    styles["Heading 1"].font.color.rgb = RGBColor(28, 55, 78)
    styles["Heading 2"].font.color.rgb = RGBColor(47, 84, 110)
    add_header_footer(doc)

    # Title page
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.space_after = Pt(8)
    r = p.add_run(TITLE)
    r.bold = True
    r.font.size = Pt(21)
    r.font.color.rgb = RGBColor(28, 55, 78)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(REPORT_TITLE)
    r.bold = True
    r.font.size = Pt(15)
    r.font.color.rgb = RGBColor(63, 94, 116)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Research progress report | Reproducible baseline predictive analysis")
    r.font.size = Pt(11)
    r.font.color.rgb = RGBColor(90, 90, 90)

    doc.add_paragraph()
    add_table(doc, ["Item", "Current state"], [
        ("Analytical period", ANALYTICAL_PERIOD),
        ("Final V2 panel", "1,615 station-month observations"),
        ("Stations", "35"),
        ("Primary target", "CPCB PM₂.₅ monthly outcome"),
        ("Primary split", "Year × Month stratified 80:20 holdout"),
        ("Random seed", str(SEED)),
        ("Primary baseline", "LightGBM"),
    ])
    doc.add_page_break()

    # Executive summary
    doc.add_heading("Executive Summary", level=1)
    doc.add_paragraph(
        "This update documents the transition from the initial ML-ready station-month dataset to a frozen "
        "multimodal V2 modeling panel, the diagnosis of the first two train/test designs, and the current baseline "
        "predictive results. The aim at this stage is deliberately limited: establish a defensible predictive baseline "
        "and a reproducible evaluation design before the later causal analysis."
    )
    doc.add_paragraph(
        "The final V2 panel contains 1,615 station-month observations from 35 stations over 2022–2025. The dataset "
        "combines the existing satellite-derived vegetation, land-surface and atmospheric features with ERA5 meteorology, "
        "WorldPop population-density features, OpenStreetMap road-infrastructure features, and a 2021 ESA WorldCover "
        "structural land-cover baseline."
    )
    doc.add_paragraph(
        "Three broad lessons emerge. First, the original holdout design was not suitable for publication-quality test "
        "evaluation because its test partition was heavily concentrated in October and November and had a very different "
        "PM₂.₅ distribution from training. Second, the data integration itself was internally consistent: row counts, "
        "keys and feature completeness were preserved. Third, after redesigning the primary split around Year × Month "
        "strata, the nonlinear tree models became the strongest baselines."
    )
    add_bullets(doc, [
        "LightGBM achieved test R² = 0.8975, RMSE = 20.38 and MAE = 10.46.",
        "Random Forest achieved test R² = 0.8759, RMSE = 22.42 and MAE = 12.24.",
        "Linear Regression and Ridge produced negative held-out R² and remain useful as diagnostic linear benchmarks.",
        "The baseline results are predictive rather than causal; feature importance is not a treatment effect.",
    ])

    # Context
    doc.add_heading("1. Research Context and Scope of Update", level=1)
    doc.add_paragraph(
        "The broader research question is whether urban green cover is associated with meaningful reductions in PM₂.₅ "
        "across Delhi NCR and whether those relationships exhibit spatial heterogeneity or thresholds. The project is "
        "designed as a multimodal spatial causal framework. This update covers only the predictive foundation needed before "
        "the causal stage."
    )
    add_bullets(doc, [
        "Freeze and audit the multimodal V2 station-month panel without altering the original V1 dataset.",
        "Document the nine CPCB stations present in the hourly master but absent from the ML-ready panel.",
        "Integrate and validate meteorological, population, road-network and structural land-cover features.",
        "Compare the split/modeling attempts and diagnose why the first holdout failed as a representative test design.",
        "Establish a reproducible Year × Month stratified 80:20 holdout and evaluate four baseline regressors.",
        "Report year-wise, season-wise, spatial-grouped and temporal-grouped diagnostics without conflating their estimands.",
    ])
    doc.add_paragraph(
        "This update does not claim a causal effect of vegetation. Feature importance is treated as predictive attribution, "
        "not a treatment effect. Likewise, the primary holdout is a within-domain interpolation-style evaluation because the "
        "same stations may contribute observations to both train and test. Spatial transferability is assessed separately "
        "through grouped validation."
    )

    # Data construction
    doc.add_heading("2. Data Construction and Station Audit", level=1)
    doc.add_heading("2.1 CPCB source and panel formation", level=2)
    doc.add_paragraph(
        "The CPCB hourly master used in the audit contains 1,377,696 hourly records across 44 stations. Duplicate checks "
        "reported zero exact duplicate rows, zero duplicate station–Timestamp combinations and zero duplicate station–year–Timestamp "
        "combinations. The ML-ready panel subsequently contains 1,615 station-month observations across 35 stations."
    )
    add_table(doc, ["Stage", "Rows", "Stations", "Interpretation"], [
        ("CPCB hourly master", "1,377,696", "44", "Source-level hourly observations after duplicate checks"),
        ("ML-ready V1", "1,615", "35", "Station-month analytical panel used as the base for V2"),
        ("V2 integrated panel", "1,615", "35", "V1 plus multimodal environmental and urban covariates"),
    ])
    fig = find_figure(root, ["data_reduction", "data-volume", "reduction"], used_figures)
    if fig:
        used_figures.add(fig.resolve().as_posix())
    add_figure(doc, fig, "Figure 1. Data-volume transition from the hourly CPCB source to the final station-month analytical panel.",
               "[Data-volume transition figure would be inserted here when present in reports/figures/modeling/]")

    doc.add_heading("2.2 Stations excluded before final ML-ready panel", level=2)
    add_table(doc, ["Category", "Stations", "Decision"], [
        ("Strong CPCB coverage; candidate recovery", "CPRI_Mathura_Road; NSIT_Dwarka; Pusa", "Investigate only if the complete Sentinel-2 / Sentinel-5P / MODIS stack can be regenerated using the same methodology."),
        ("Weak/irregular coverage", "Lodhi_Road", "Exclude from primary V2; possible future sensitivity analysis under an explicit quality rule."),
        ("No usable PM₂.₅ outcome", "Commonwealth_Sports_Complex; IGNOU_Maidan_Garhi; JNU; NSUT_Jaffarpur; Talkatora_Garden", "Exclude. The audited 2025 records had zero valid PM₂.₅."),
    ])
    doc.add_paragraph(
        "The conservative choice was to keep the existing 35-station panel rather than expand it midstream. This preserves "
        "methodological consistency and avoids introducing station-specific feature-generation decisions after the satellite/urban "
        "feature architecture was already frozen. The three stations with strong CPCB coverage remain candidates for later recovery, "
        "but were not retroactively inserted into V2."
    )
    doc.add_heading("2.3 IIT Delhi singleton", level=2)
    doc.add_paragraph(
        "IIT_Delhi contributes one valid station-month observation (2025-12) in the current V2 panel. It was retained because V2 "
        "was frozen, but assigned train-only in the primary split because a single observation cannot be meaningfully partitioned "
        "between training and testing. It is not used as an independent test location and should not drive claims of station-level generalization."
    )

    # Feature integration
    doc.add_heading("3. Multimodal Feature Integration", level=1)
    add_table(doc, ["Feature layer", "Source / method", "Role in V2", "Key limitation"], [
        ("Vegetation / optical", "Sentinel-2 indices and multi-scale summaries", "Dynamic vegetation / surface-condition observations", "Dependent on valid observations and cloud filtering"),
        ("Atmospheric composition proxy", "Sentinel-5P NO₂", "Anthropogenic atmospheric context", "NO₂ is a proxy for source activity, not a direct PM₂.₅ measurement"),
        ("MODIS summaries", "MODIS NDVI, EVI and LST", "Multi-scale vegetation and thermal context", "Coarser spatial scale than Sentinel-2"),
        ("Meteorology", "ERA5-Land + ERA5 boundary-layer height", "Short-term physical drivers of PM₂.₅", "Reanalysis rather than local station observations"),
        ("Population", "WorldPop density summaries", "Urban intensity / exposure context", "Static population representation in the extracted panel"),
        ("Road infrastructure", "OSM density at 100/250/500/1000 m + major-road density", "Structural transport intensity", "January 2025 snapshot used as a static infrastructure proxy"),
        ("Structural land cover", "ESA WorldCover 2021 fractional composition", "Pre-study contextual land-use baseline", "Does not represent 2022–2025 land-cover change"),
    ])
    doc.add_heading("3.1 OpenStreetMap audit", level=2)
    doc.add_paragraph(
        "The OSM extraction was technically successful after moving away from station-by-station Overpass queries to a local PBF "
        "workflow using PyOsmium. The validated feature table contains 1,615 rows, 35 stations, zero missing feature values and zero "
        "duplicate station-month keys. The extraction reported 1,600,339 highway geometries and 58,713 major-road geometries, with metric "
        "calculations performed in EPSG:32643."
    )
    doc.add_paragraph(
        "The main limitation is interpretive rather than computational: the road layer is a January 2025 OSM snapshot applied across "
        "a 2022–2025 panel. It should therefore be described as a static structural infrastructure proxy, not as year-specific traffic intensity."
    )
    fig = find_figure(root, ["feature_group", "group_importance", "contribution"], used_figures)
    if fig:
        used_figures.add(fig.resolve().as_posix())
    add_figure(doc, fig, "Figure 2. Feature-group contribution reported by the modeling workflow. These values are predictive importance shares, not causal effects.",
               "[Feature-group contribution figure would be inserted here when present in reports/figures/modeling/]")

    doc.add_heading("3.2 WorldCover static baseline", level=2)
    doc.add_paragraph(
        "ESA WorldCover v200 (2021) was incorporated as a categorical land-cover baseline using class fractions rather than arithmetic "
        "raster averages. The extracted features were fractional composition measures such as built, grass, cropland and water within "
        "the specified buffers. Validation reported 1,615 rows, 35 stations and no missing values."
    )
    doc.add_paragraph(
        "The scientific rationale for using a 2021 static map is that it represents a pre-study structural context rather than a time-varying "
        "vegetation exposure. This is useful when the later causal design must distinguish background land-use structure from the dynamic green-cover treatment."
    )
    doc.add_paragraph(
        "The limitation is explicit: land-cover changes after 2021 are not represented."
    )

    # Split design
    doc.add_heading("4. Split Design: Three Attempts and Why the First Two Were Reconsidered", level=1)
    doc.add_heading("4.1 Attempt 1 — initial holdout", level=2)
    doc.add_paragraph(
        "The first modeling split produced 1,344 training observations and 271 test observations. Although several hard integrity checks passed, "
        "the test set was heavily concentrated in October and November. The test set contained 261 post-monsoon observations and 56 winter observations, "
        "whereas training contained only six post-monsoon observations and 349 winter observations. This produced a large target-distribution shift: "
        "training mean PM₂.₅ was approximately 85.21 while test mean PM₂.₅ was approximately 164.67. LightGBM had very high training fit but negative held-out R²."
    )
    doc.add_paragraph(
        "The main interpretation was not that the algorithm was inherently incapable of modeling PM₂.₅, but that the evaluation design exposed the model "
        "to a test regime scarcely represented in training."
    )
    fig = find_figure(root, ["month", "split", "diagnostic", "temporal"], used_figures)
    if fig:
        used_figures.add(fig.resolve().as_posix())
    add_figure(doc, fig, "Figure 3. Month-wise preservation of temporal structure in the redesigned primary split.",
               "[Month-wise split diagnostic would be inserted here when present in reports/figures/modeling/]")

    doc.add_heading("4.2 Attempt 2 — independent audit and alternative month/year split", level=2)
    doc.add_paragraph(
        "The independent audit shared for this work confirmed the above diagnosis and evaluated an alternative month/year-aware split without overwriting the locked initial outputs. "
        "That experiment produced a train/test target-mean difference of only about 2.26 PM₂.₅ units and reported LightGBM test R² ≈ 0.958 and Random Forest test R² ≈ 0.920. "
        "The result was useful as a methodological check because it demonstrated that the initial negative R² values were strongly affected by split composition."
    )
    doc.add_paragraph(
        "Those alternative results were not simply adopted as the final published numbers. Instead, they informed the independently generated V2 split so that the final modeling workflow remained reproducible from the canonical V2 dataset."
    )

    doc.add_heading("4.3 Attempt 3 — current primary split", level=2)
    doc.add_paragraph(
        "The present split uses an exact 80:20 test fraction (1,292 train; 323 test) with Year × Month stratification and largest-remainder allocation. "
        "The singleton IIT_Delhi observation is fixed in training. The V2 dataset itself is treated as read-only and the split is validated both in memory and after writing to disk."
    )
    add_table(doc, ["Constraint", "Current status"], [
        ("Total observations", "1,615"),
        ("Training / testing", "1,292 / 323"),
        ("Exact ratio", "80.0% / 20.0%"),
        ("Years represented in both sets", "2022, 2023, 2024, 2025"),
        ("Every eligible year-month stratum has test coverage", "Yes"),
        ("Train/test key overlap", "0"),
        ("Row universe preserved", "Yes"),
        ("IIT_Delhi", "Train-only; n=1"),
    ])
    fig = find_figure(root, ["holdout", "split_size", "train", "test"], used_figures)
    if fig:
        used_figures.add(fig.resolve().as_posix())
    add_figure(doc, fig, "Figure 4. Final primary holdout size: exact 80:20 split.",
               "[Final holdout-size figure would be inserted here when present in reports/figures/modeling/]")

    # Baseline results
    doc.add_heading("5. Baseline Modeling Results", level=1)
    doc.add_paragraph(
        "Four regressors were evaluated: Linear Regression, Ridge Regression, Random Forest and LightGBM. The reported cross-validation scores were computed on the training partition only. "
        "The held-out test set was kept separate for final evaluation."
    )
    results = get_model_results(root)
    add_table(doc, ["Model", "Train R²", "Test R²", "Test RMSE", "Test MAE", "Median AE", "CV R² mean", "CV R² SD", "CV RMSE", "R² gap"], [
        (r["model"], fmt(r["train_r2"]), fmt(r["test_r2"]), fmt(r["test_rmse"]), fmt(r["test_mae"]), fmt(r["test_median_ae"]), fmt(r["cv_r2_mean"]), fmt(r["cv_r2_std"], 4), fmt(r["cv_rmse_mean"]), fmt(r["r2_gap"]))
        for r in results
    ])
    fig = find_figure(root, ["model_comparison", "baseline_model", "comparison"], used_figures)
    if fig:
        used_figures.add(fig.resolve().as_posix())
    add_figure(doc, fig, "Figure 5. Held-out test comparison across the four baseline regressors.",
               "[Model-comparison figure would be inserted here when present in reports/figures/modeling/]")

    doc.add_heading("5.1 Best baseline model: LightGBM", level=2)
    doc.add_paragraph(
        "LightGBM is the most suitable primary baseline among the four models. The reason is not only its highest test R². It also has the lowest test RMSE and MAE, the strongest cross-validation mean R², "
        "a very small CV standard deviation, and a substantially smaller train-to-test R² gap than the linear models. Its test R² of 0.8975 remains reasonably close to its training-only CV R² mean of 0.9470."
    )
    doc.add_paragraph(
        "Random Forest is a strong secondary baseline. Its test R² of 0.8759 and MAE of 12.24 remain good, and its train/test gap is only slightly larger than LightGBM. Keeping both tree ensembles provides robustness evidence that the nonlinear result is not dependent on a single boosting implementation."
    )
    doc.add_paragraph(
        "Linear Regression and Ridge are retained as benchmark models rather than discarded. Their negative held-out R² values indicate that the current multimodal predictor space is not adequately represented by a simple linear specification under the current split."
    )
    fig = find_figure(root, ["cv", "held", "cross", "validation"], used_figures)
    if fig:
        used_figures.add(fig.resolve().as_posix())
    add_figure(doc, fig, "Figure 6. Relationship between five-fold training CV R² and held-out test R².",
               "[CV versus test figure would be inserted here when present in reports/figures/modeling/]")

    # Year and season
    doc.add_heading("5.2 Year-wise behavior", level=2)
    add_table(doc, ["Model", "2022 R²", "2023 R²", "2024 R²", "2025 R²"], [
        (m, *(f"{dict((year, r2) for r2, year in vals)[y]:.3f}" for y in [2022, 2023, 2024, 2025])) for m, vals in YEARWISE.items()
    ])
    doc.add_paragraph(
        "Both nonlinear models remain positive across all four years. The weakest year for both is 2024, which is worth carrying forward into residual and data-quality diagnostics rather than interpreting immediately as a scientific effect. "
        "The consistency across 2022, 2023 and 2025 is encouraging, while the 2024 decline suggests a year-specific regime or data characteristic deserves further inspection."
    )
    fig = find_figure(root, ["yearwise", "year", "held-out"], used_figures)
    if fig:
        used_figures.add(fig.resolve().as_posix())
    add_figure(doc, fig, "Figure 7. Year-wise held-out performance for Random Forest and LightGBM.",
               "[Year-wise performance figure would be inserted here when present in reports/figures/modeling/]")

    doc.add_heading("5.3 Season-wise behavior", level=2)
    add_table(doc, ["Model", "Winter R²", "Summer R²", "Monsoon R²", "Post-monsoon R²"], [
        (m, f"{v['Winter']:.3f}", f"{v['Summer']:.3f}", f"{v['Monsoon']:.3f}", f"{v['Post-monsoon']:.3f}") for m, v in SEASONWISE.items()
    ])
    doc.add_paragraph(
        "The tree-based models remain positive in all four seasons, with strongest performance in winter and weakest in monsoon. The linear models fail dramatically in monsoon because a relatively small number of extreme residuals makes the squared-error criterion unstable. "
        "This seasonal pattern is best treated as a robustness finding and a prompt for residual/meterological inspection, not as proof of a biological or causal mechanism."
    )
    fig = find_figure(root, ["seasonwise", "season", "performance"], used_figures)
    if fig:
        used_figures.add(fig.resolve().as_posix())
    add_figure(doc, fig, "Figure 8. Season-wise held-out performance for Random Forest and LightGBM.",
               "[Season-wise performance figure would be inserted here when present in reports/figures/modeling/]")

    # grouped validation
    doc.add_heading("6. Secondary Generalization Checks", level=1)
    doc.add_paragraph(
        "The primary holdout is not a spatial generalization test because the same monitoring stations can contribute observations to both train and test. "
        "Secondary grouped validation was therefore reported separately. A station-grouped evaluation is closer to the question of transfer to previously unseen monitoring locations, while year-grouped validation is a temporal transfer diagnostic."
    )
    add_table(doc, ["Model", "Spatial grouped CV mean R²", "Temporal grouped CV mean R²"], [
        (m, f"{sp:.4f}", f"{tp:.4f}") for m, sp, tp in GROUPED_RESULTS
    ])
    doc.add_paragraph(
        "Grouped results support retaining both nonlinear models in the later analysis. LightGBM remains strong under both spatial and temporal grouping, while Random Forest is slightly stronger on the spatial grouped metric. These scores should not be pooled with the primary holdout because they answer different generalization questions."
    )
    fig = find_figure(root, ["grouped", "spatial", "temporal", "generalization"], used_figures)
    if fig:
        used_figures.add(fig.resolve().as_posix())
    add_figure(doc, fig, "Figure 9. Secondary spatial and temporal generalization checks.",
               "[Grouped-validation figure would be inserted here when present in reports/figures/modeling/]")

    # residuals
    doc.add_heading("6.1 Residual diagnostics", level=2)
    add_table(doc, ["Residual statistic", "Value"], [
        ("Mean residual", RESIDUALS["mean"]),
        ("Median residual", RESIDUALS["median"]),
        ("Residual SD", RESIDUALS["std"]),
        ("Residual skewness", RESIDUALS["skewness"]),
        ("Maximum underprediction", RESIDUALS["max_underprediction"]),
        ("Maximum overprediction", RESIDUALS["max_overprediction"]),
    ])
    doc.add_paragraph(
        "The residual summary indicates a near-zero mean residual and median, but strong negative skewness. The maximum reported overprediction of approximately 262.7 PM₂.₅ units is driven by at least one unusual observation/model mismatch. "
        "This should be investigated at the original data level before any substantive interpretation."
    )
    fig = find_figure(root, ["residual", "error"], used_figures)
    if fig:
        used_figures.add(fig.resolve().as_posix())
    add_figure(doc, fig, "Figure 10. Residual summary for the current LightGBM evaluation.",
               "[Residual-summary figure would be inserted here when present in reports/figures/modeling/]")

    # Findings
    doc.add_heading("7. What the Current Modeling Results Support", level=1)
    add_bullets(doc, [
        "Nonlinearity: the tree ensembles substantially outperform the linear specifications on the held-out panel.",
        "Multimodal usefulness: combining vegetation, meteorology, atmospheric proxies and urban-structure variables yields a strong predictive baseline.",
        "Temporal robustness: the nonlinear models remain positive across all four years, with a weaker 2024 regime requiring follow-up.",
        "Spatial robustness: grouped validation remains substantially positive for the nonlinear models, although this should not be equated with universal spatial transferability.",
        "Model choice: LightGBM is the most defensible primary predictive baseline, while Random Forest is the main robustness comparator.",
    ])
    doc.add_heading("7.1 What the results do not establish", level=2)
    add_bullets(doc, [
        "They do not establish a causal green-cover effect.",
        "Tree-model feature importance is not an elasticity, treatment effect or causal contribution.",
        "They do not establish that the model will work in a new city, an entirely unmonitored Delhi location or a future year beyond the observed period.",
    ])

    doc.add_heading("7.2 Feature-group findings", level=2)
    models = ["Linear Regression", "Ridge Regression", "Random Forest", "LightGBM"]
    add_table(doc, ["Feature group", *models], [
        (g, *(f"{vals[m]:.2f}%" for m in models)) for g, vals in FEATURE_GROUPS.items()
    ])
    doc.add_paragraph(
        "The LightGBM workflow assigns a large share of predictive importance to green-cover features (46.06%) while retaining substantial meteorological (23.07%) and pollution/anthropogenic (13.47%) contributions. "
        "That is encouraging for the research question, but the result remains predictive evidence. The very large spatial/temporal share reported by Random Forest is a reminder that calendar and location structure can dominate a model's explanatory power."
    )
    fig = find_figure(root, ["feature_group", "importance", "contribution"])
    add_figure(doc, fig, "Figure 11. Feature-group importance summary from the baseline modeling workflow.",
               "[Feature-group importance figure would be inserted here when present in reports/figures/modeling/]")

    # Reproducibility
    doc.add_heading("8. Reproducibility and File Organization", level=1)
    doc.add_paragraph(
        "The workflow is intended to be reproducible from the frozen V2 dataset. The canonical dataset is treated as read-only; splitting creates derivative train/test files; modeling reads those derivatives; reports and figures are written separately. The original V1 and V2 datasets are not reconstructed from model outputs."
    )
    add_table(doc, ["Artifact", "Purpose"], [
        ("data/ml_ready/master_modeling_dataset.csv", "Original V1 ML-ready station-month panel; retained as historical baseline."),
        ("data/ml_ready/master_modeling_dataset_v2.csv", "Canonical multimodal V2 panel used as the source for final modeling splits."),
        ("data/modeling_final/train.csv", "Primary training partition (1,292 rows)."),
        ("data/modeling_final/test.csv", "Primary held-out test partition (323 rows)."),
        ("data/modeling_final/split_manifest.csv", "Split provenance: seed, row counts, station handling and allocation logic."),
        ("data/modeling_final/validation_report.csv", "Pre- and post-split integrity checks."),
        ("data/modeling_final/distribution_diagnostics.csv", "Year, month, year-month, season and station allocation diagnostics."),
        ("data/03_features/feat_era5_met.csv", "ERA5 / ERA5-Land station-month meteorological features."),
        ("data/03_features/feat_worldpop.csv", "Population-density features aligned to the station-month panel."),
        ("data/03_features/feat_osm_roads.csv", "OSM structural road-density features."),
        ("data/03_features/feat_worldcover.csv", "Static 2021 land-cover fractions."),
    ])
    doc.add_paragraph(
        "The primary split is deterministic with seed 42. The eligible panel is stratified by calendar year × month, with largest-remainder allocation used to obtain exactly 323 test observations. IIT_Delhi is explicitly held in training because its group contains one observation. After sampling, the workflow verifies key uniqueness, zero train/test overlap, exact row-universe preservation, year coverage, year-month test coverage and schema preservation."
    )
    fig = find_figure(root, ["distribution", "month", "diagnostic"], used_figures)
    if fig:
        used_figures.add(fig.resolve().as_posix())
    add_figure(doc, fig, "Figure 12. Month-wise split diagnostics reproduced from the finalized allocation.",
               "[Distribution diagnostics figure would be inserted here when present in reports/figures/modeling/]")

    doc.add_heading("8.1 Environment and extraction assumptions", level=2)
    add_bullets(doc, [
        "OSM extraction used a local PBF workflow with PyOsmium after repeated Overpass / OSMnx timeouts.",
        "ERA5 monthly variables were computed server-side in Google Earth Engine to avoid hourly data-volume bottlenecks.",
        "WorldCover was extracted as categorical fractions from ESA WorldCover v200 (2021).",
        "The January 2025 OSM layer is treated as a static structural infrastructure proxy across 2022–2025.",
        "ESA WorldCover 2021 is treated as a static pre-study land-cover baseline.",
        "Population density is treated as an urban-intensity feature, not as a causal treatment.",
        "Green-cover predictor importance is interpreted as predictive contribution only.",
    ])

    # Issues
    doc.add_heading("9. Issues Flagged for Follow-up", level=1)
    add_table(doc, ["Issue", "Evidence", "Action"], [
        ("CPCB anomalous value", "Chandni_Chowk, Nov 2024, PM₂.₅ = 5.0 while model predictions are around 200–270", "Inspect the original CPCB record and quality flags before substantive interpretation."),
        ("Extreme linear extrapolation", "Bawana, Jul 2022, linear predictions > 1,800 for observed value near 43.5", "Keep linear models as diagnostics; audit feature scaling and outlier sensitivity if discussed."),
        ("2024 weaker nonlinear performance", "LightGBM R² = 0.740 and RF R² = 0.718", "Inspect year-specific feature distributions, missingness and station coverage."),
        ("Monsoon lower performance", "LightGBM R² = 0.529; RF R² = 0.444", "Treat as robustness finding; inspect residuals and meteorological regime."),
        ("Static OSM baseline", "January 2025 infrastructure used for 2022–2025", "State clearly that it represents structural road layout, not historical annual traffic volume."),
        ("Static WorldCover baseline", "2021 product used for 2022–2025", "State clearly that land-cover change after 2021 is not represented."),
        ("Singleton station", "IIT_Delhi n=1", "Keep train-only and exclude from station-level generalization claims."),
        ("Candidate recoveries", "CPRI Mathura Road, NSIT Dwarka, Pusa have strong CPCB coverage but were absent from V2", "Consider later only if the full multimodal feature stack can be generated consistently."),
    ])
    doc.add_paragraph(
        "No further data correction was made before this baseline because the working principle was to freeze V2 before extensive tuning. "
        "This prevents an iterative loop in which the dataset is repeatedly altered until predictive scores improve. The present results therefore serve as a clean baseline against which later, explicitly justified changes can be compared."
    )

    # Next stage and closing
    doc.add_heading("10. Recommended Next Stage", level=1)
    add_bullets(doc, [
        "Conduct a focused audit of influential observations, especially the 2024 Chandni_Chowk PM₂.₅ value of 5.0 and the extreme linear predictions.",
        "Freeze the current LightGBM and Random Forest baseline metrics as the reference predictive results for Update 2.",
        "Use station-grouped validation as the principal secondary spatial generalization test, keeping it separate from the primary balanced holdout.",
        "Examine green-cover features at 100, 250, 500 and 1000 m for consistency and possible multiscale threshold behavior.",
        "Use the predictive findings to define the treatment, confounder and nuisance-variable roles for the planned Double Machine Learning / Causal Forest analysis.",
        "Only after these checks consider feature reduction, hyperparameter optimization or additional station recovery.",
    ])

    # Appendix: automatically include any additional modeling PNGs not already used above.
    fig_dir = root / "reports" / "figures" / "modeling"
    if fig_dir.exists():
        remaining = [p for p in sorted(fig_dir.glob("*.png")) if p.resolve().as_posix() not in used_figures]
        if remaining:
            doc.add_heading("Appendix A. Additional Final-V2 Modeling Visualizations", level=1)
            doc.add_paragraph(
                "The following figures were present in reports/figures/modeling at generation time but were not assigned to a named section above. "
                "They are included here so that the report remains reproducible with the visualization outputs generated by the final V2 modeling workflow."
            )
            for idx, fig_path in enumerate(remaining, 1):
                try:
                    doc.add_picture(str(fig_path), width=Inches(6.4))
                    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
                    add_caption(doc, f"Appendix Figure A{idx}. {fig_path.stem.replace('_', ' ').replace('-', ' ').title()}.")
                except Exception as exc:
                    doc.add_paragraph(f"Could not embed {fig_path.name}: {exc}")

    doc.add_heading("11. Closing Assessment", level=1)
    doc.add_paragraph(
        "At this stage, the baseline modeling pipeline is in a materially better position than at the start of the modeling work. The key improvement was not simply obtaining a higher R²; it was recognizing that the evaluation split is part of the scientific design. "
        "After resplitting the frozen V2 panel to preserve calendar structure, the nonlinear models produced stable positive test performance across the study period, while linear models remained poor."
    )
    p = doc.add_paragraph()
    r = p.add_run(
        "Current conclusion: LightGBM is the best baseline predictive model for the current multimodal Delhi NCR PM₂.₅ panel, "
        "with Random Forest providing an independent nonlinear comparator. The evidence is sufficiently strong to proceed to "
        "deeper spatial diagnostics and causal-model preparation, while keeping the stated data and evaluation limitations visible in the research record."
    )
    r.bold = True

    # Small generation note
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Generated reproducibly by scripts/generate_baseline_data_modelling_update.py")
    r.font.size = Pt(8)
    r.font.color.rgb = RGBColor(120, 120, 120)

    doc.save(output)


def main() -> None:
    parser = argparse.ArgumentParser(description="Generate Update 2 baseline modeling research report as DOCX")
    parser.add_argument("--root", type=Path, default=Path.cwd(), help="Project root directory")
    parser.add_argument(
        "--output",
        type=Path,
        default=Path("reports") / "baseline_data_modelling_update.docx",
        help="Output DOCX path relative to --root unless absolute",
    )
    args = parser.parse_args()
    root = args.root.resolve()
    output = args.output if args.output.is_absolute() else root / args.output
    build_document(root, output)
    print(f"Generated: {output}")


if __name__ == "__main__":
    main()